"""Generate realistic, professional PDF + DOCX reference documents for the workshop.

Outputs (written to ../Fraud_Analytics/docs, path passed as argv[1]):
  - fraud_event_report.pdf        (SIU investigation report — letterhead, exhibits, sign-off)
  - eu_compliance_policy.pdf      (EU/GDPR fraud analytics compliance framework)
  - genie_questions.docx          (10+ Genie questions)
"""
import sys
import os

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, ListFlowable, ListItem,
)

NAVY = colors.HexColor("#14243d")
BLUE = colors.HexColor("#2f6df6")
RED = colors.HexColor("#c0392b")
GREY = colors.HexColor("#6b7280")
LINE = colors.HexColor("#d7dee8")
LIGHT = colors.HexColor("#f2f5fa")

OUT = sys.argv[1] if len(sys.argv) > 1 else "."


# ---------------------------------------------------------------- styles
def styles():
    s = getSampleStyleSheet()
    s.add(ParagraphStyle("Cls", fontName="Helvetica-Bold", fontSize=7.5, textColor=RED,
                         spaceAfter=2, alignment=TA_LEFT))
    s.add(ParagraphStyle("DocTitle", fontName="Helvetica-Bold", fontSize=19, textColor=NAVY,
                         spaceBefore=6, spaceAfter=4, leading=23))
    s.add(ParagraphStyle("SubTitle", fontName="Helvetica", fontSize=11, textColor=GREY,
                         spaceAfter=12, leading=15))
    s.add(ParagraphStyle("H1", fontName="Helvetica-Bold", fontSize=12.5, textColor=NAVY,
                         spaceBefore=14, spaceAfter=6, leading=16))
    s.add(ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=10.5, textColor=colors.HexColor("#1b3b6f"),
                         spaceBefore=9, spaceAfter=4, leading=14))
    s.add(ParagraphStyle("Body2", fontName="Helvetica", fontSize=9.7, textColor=colors.HexColor("#1f2937"),
                         spaceAfter=7, leading=14.5, alignment=TA_JUSTIFY))
    s.add(ParagraphStyle("Bullet2", parent=s["Body2"], leftIndent=12, spaceAfter=3, alignment=TA_LEFT))
    s.add(ParagraphStyle("Small", fontName="Helvetica", fontSize=8, textColor=GREY, leading=11))
    s.add(ParagraphStyle("Cell", fontName="Helvetica", fontSize=8.5, textColor=colors.HexColor("#1f2937"), leading=11))
    s.add(ParagraphStyle("CellH", fontName="Helvetica-Bold", fontSize=8.5, textColor=colors.white, leading=11))
    return s


S = styles()


def make_doc(path, brand, unit, classification):
    doc = BaseDocTemplate(path, pagesize=A4,
                          leftMargin=2.0 * cm, rightMargin=2.0 * cm,
                          topMargin=2.4 * cm, bottomMargin=1.9 * cm,
                          title=brand)

    def header_footer(canvas, d):
        canvas.saveState()
        # header band
        canvas.setFillColor(NAVY)
        canvas.rect(0, A4[1] - 1.5 * cm, A4[0], 1.5 * cm, fill=1, stroke=0)
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 11)
        canvas.drawString(2.0 * cm, A4[1] - 1.0 * cm, brand)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#aeb9cc"))
        canvas.drawRightString(A4[0] - 2.0 * cm, A4[1] - 1.0 * cm, unit)
        # footer
        canvas.setStrokeColor(LINE)
        canvas.setLineWidth(0.5)
        canvas.line(2.0 * cm, 1.5 * cm, A4[0] - 2.0 * cm, 1.5 * cm)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(RED)
        canvas.drawString(2.0 * cm, 1.1 * cm, classification)
        canvas.setFillColor(GREY)
        canvas.drawCentredString(A4[0] / 2, 1.1 * cm, brand)
        canvas.drawRightString(A4[0] - 2.0 * cm, 1.1 * cm, "Page %d" % d.page)
        canvas.restoreState()

    frame = Frame(doc.leftMargin, doc.bottomMargin,
                  doc.width, doc.height - 0.4 * cm, id="main")
    doc.addPageTemplates([PageTemplate(id="std", frames=[frame], onPage=header_footer)])
    return doc


def control_table(rows):
    data = [[Paragraph(k, S["Cell"]), Paragraph(v, S["Cell"])] for k, v in rows]
    t = Table(data, colWidths=[4.2 * cm, 12.3 * cm])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, LINE),
        ("BACKGROUND", (0, 0), (0, -1), LIGHT),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def data_table(header, rows, widths):
    data = [[Paragraph(h, S["CellH"]) for h in header]]
    for r in rows:
        data.append([Paragraph(str(c), S["Cell"]) for c in r])
    t = Table(data, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("GRID", (0, 0), (-1, -1), 0.5, LINE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
    ]))
    return t


def bullets(items):
    return ListFlowable(
        [ListItem(Paragraph(t, S["Bullet2"]), leftIndent=10, value="•") for t in items],
        bulletType="bullet", start="•", leftIndent=8,
    )


def hr():
    return HRFlowable(width="100%", thickness=0.6, color=LINE, spaceBefore=4, spaceAfter=8)


# ================================================================ FRAUD REPORT
def build_fraud_report():
    path = os.path.join(OUT, "fraud_event_report.pdf")
    doc = make_doc(path, "Allianz Insurance plc",
                   "Special Investigations Unit", "RESTRICTED — SIU EYES ONLY")
    e = []
    e.append(Spacer(1, 6))
    e.append(Paragraph("RESTRICTED — SPECIAL INVESTIGATIONS UNIT", S["Cls"]))
    e.append(Paragraph("Organised Claims Fraud — Investigation Report", S["DocTitle"]))
    e.append(Paragraph("Suspected organised motor &amp; home claims fraud ring — "
                       "London and North West regions", S["SubTitle"]))
    e.append(control_table([
        ("Report reference", "SIU-2024-0417"),
        ("Classification", "Restricted — SIU Eyes Only"),
        ("Prepared by", "J. Harrington, Senior Investigator, Special Investigations Unit"),
        ("Reviewed by", "D. Osei, Head of Claims Fraud"),
        ("Data &amp; analytics", "Claims Intelligence &amp; Fraud Analytics team"),
        ("Reporting period", "01 January 2023 – 31 December 2024"),
        ("Date of issue", "17 January 2025"),
        ("Distribution", "Head of Claims; Group Financial Crime; Legal &amp; Compliance"),
        ("Status", "Referred to financial-crime pathway — 3 claims denied, 2 under review"),
    ]))
    e.append(Spacer(1, 8))
    e.append(Paragraph("This report concerns fictitious data prepared for an analytics "
                       "workshop. All names, policy numbers and claim identifiers are "
                       "synthetic. Figures reconcile to the governed "
                       "<b>gold_fraud_claims</b> dataset.", S["Small"]))

    e.append(Paragraph("1.  Executive Summary", S["H1"]))
    e.append(hr())
    e.append(Paragraph(
        "Between January 2023 and December 2024 the Special Investigations Unit (SIU) "
        "identified a cluster of 41 claims exhibiting a consistent and statistically "
        "improbable pattern. The cluster spans <b>Auto</b> and <b>Home</b> policy lines, is "
        "concentrated in the <b>London</b> and <b>North West</b> regions, and is "
        "overwhelmingly submitted through the <b>Online</b> and <b>Mobile App</b> channels. "
        "Individually the claims are unremarkable; in combination they drive the internal "
        "<b>fraud_risk_score</b> to 3 or 4 on a 0–4 scale.", S["Body2"]))
    e.append(Paragraph("The recurring four-signal fingerprint is:", S["Body2"]))
    e.append(bullets([
        "<b>High claim value</b> — payout materially above the policy-type average "
        "(<b>high_value_flag</b>, triggered above £20,000).",
        "<b>Fast reporting</b> — incident reported within 0–2 days "
        "(<b>fast_report_flag</b>, report_lag_days ≤ 2).",
        "<b>New customer tenure</b> — policy opened under twelve months before the loss "
        "(<b>new_customer_flag</b>, tenure_years &lt; 1).",
        "<b>Low credit score</b> — policyholder credit score below 500 at inception "
        "(<b>low_credit_flag</b>).",
    ]))
    e.append(Paragraph(
        "Of the 41 claims, <b>29 carried a fraud_risk_score of 3 or higher</b> and 17 were "
        "confirmed fraudulent (<b>is_fraud = 1</b>). Estimated gross exposure avoided through "
        "early detection and denial is approximately <b>£612,000</b>.", S["Body2"]))

    e.append(Paragraph("2.  Background and Trigger", S["H1"]))
    e.append(hr())
    e.append(Paragraph(
        "The investigation was triggered by an automated alert from the fraud analytics "
        "pipeline. During the routine Bronze → Silver → Gold refresh, the aggregate table "
        "<b>gold_fraud_by_region</b> surfaced an anomaly: the fraud rate for Auto claims in "
        "London rose from a baseline near 7% to over 19% across two consecutive months, while "
        "claim volume rose only modestly. A rise in fraud <i>rate</i> not explained by a "
        "proportional rise in claim <i>volume</i> is a classic indicator of a coordinated "
        "push rather than organic growth.", S["Body2"]))
    e.append(Paragraph(
        "A secondary trigger was the channel mix. In the legitimate population claims are "
        "spread fairly evenly across Branch, Online, Broker, Call Center and Mobile App. "
        "Within the flagged cluster, over 70% of claims originated from Online and Mobile App "
        "— the two lowest-friction channels at first notice of loss (FNOL) — consistent with "
        "deliberate avoidance of adjuster contact.", S["Body2"]))

    e.append(Paragraph("3.  Anatomy of the Cluster", S["H1"]))
    e.append(hr())
    e.append(Paragraph("3.1  Risk-score distribution vs. the book", S["H2"]))
    e.append(Paragraph("The single most useful artefact was the fraud_risk_score. Its "
                       "distribution in the cluster differed starkly from the portfolio:", S["Body2"]))
    e.append(data_table(
        ["Risk score", "Cluster claims", "Portfolio baseline (approx.)"],
        [["0", "1", "34%"], ["1", "4", "33%"], ["2", "7", "20%"],
         ["3", "18", "9%"], ["4", "11", "4%"]],
        [5 * cm, 5 * cm, 6.5 * cm]))
    e.append(Spacer(1, 6))
    e.append(Paragraph(
        "In the normal book roughly one claim in eight scores 3 or 4. In this cluster more "
        "than seven in ten did. A score of 4 — all four indicators firing — is rare enough in "
        "the legitimate population to warrant manual review on its own.", S["Body2"]))
    e.append(Paragraph("3.2  Reporting-lag and tenure behaviour", S["H2"]))
    e.append(Paragraph(
        "The median reporting lag in the cluster was 1 day, and 33 of 41 claims were reported "
        "within 48 hours. Cross-referencing the policyholder dimension showed short tenure "
        "(new_customer_flag), implausibly high prior-claim counts for the tenure, and a median "
        "credit score of 430 versus a portfolio median near 575.", S["Body2"]))

    e.append(PageBreak())
    e.append(Paragraph("4.  Illustrative Cases (Exhibits)", S["H1"]))
    e.append(hr())
    e.append(Paragraph("<b>Exhibit A — Auto / Theft (score 4, confirmed fraud).</b> "
        "Comprehensive Auto policy opened Online in the North West. Eleven days later the "
        "insured SUV was reported stolen, filed same-day (report_lag_days = 0), value ≈ £34,000 "
        "(high_value_flag). Credit score 402 (low_credit_flag); tenure under one month "
        "(new_customer_flag). All four flags fired. The vehicle was found to have been exported "
        "before the policy was opened. <b>Outcome: claim denied, is_fraud = 1.</b>", S["Body2"]))
    e.append(Paragraph("<b>Exhibit B — Home / Burglary (score 3, confirmed fraud).</b> "
        "London contents claim of £28,500 for jewellery and electronics, reported within one "
        "day, four prior claims across two lines within eight months, and no police report "
        "filed. Three flags fired. <b>Outcome: claim denied, is_fraud = 1.</b>", S["Body2"]))
    e.append(Paragraph("<b>Exhibit C — Auto / Collision (score 2, legitimate — counter-example).</b> "
        "London collision of £22,000 triggered high_value_flag and new_customer_flag, but was "
        "reported after nine days, credit score 690, police report filed, no prior claims. "
        "Manual review cleared it. <b>Outcome: claim paid, is_fraud = 0.</b> A high value alone "
        "is not fraud; it is the <i>combination</i> of signals that matters.", S["Body2"]))

    e.append(Paragraph("5.  Schedule of Representative Claims", S["H1"]))
    e.append(hr())
    e.append(data_table(
        ["Claim", "Line / Type", "Region", "Channel", "Amount", "Lag", "Score", "Fraud"],
        [["CLM50xxx1", "Auto / Theft", "North West", "Online", "£34,100", "0", "4", "Yes"],
         ["CLM50xxx2", "Home / Burglary", "London", "Mobile App", "£28,500", "1", "3", "Yes"],
         ["CLM50xxx3", "Auto / Collision", "London", "Broker", "£22,000", "9", "2", "No"],
         ["CLM50xxx4", "Home / Accidental", "North West", "Online", "£19,800", "1", "3", "Yes"],
         ["CLM50xxx5", "Auto / Theft", "London", "Mobile App", "£31,250", "2", "4", "Yes"]],
        [2.6 * cm, 3.2 * cm, 2.4 * cm, 2.2 * cm, 2.0 * cm, 1.2 * cm, 1.3 * cm, 1.5 * cm]))
    e.append(Spacer(1, 6))

    e.append(Paragraph("6.  Financial Impact", S["H1"]))
    e.append(hr())
    e.append(bullets([
        "Claims in cluster: <b>41</b>",
        "Confirmed fraudulent (is_fraud = 1): <b>17</b>",
        "Gross payout avoided (denied claims): <b>≈ £612,000</b>",
        "Under review at time of writing: <b>2</b> (≈ £47,000 combined exposure)",
        "False-positive rate on manual review: <b>~19%</b> (flagged but cleared)",
    ]))

    e.append(Paragraph("7.  Recommendations", S["H1"]))
    e.append(hr())
    e.append(bullets([
        "Elevate all fraud_risk_score = 4 claims to mandatory manual review before payment.",
        "Add FNOL friction for Online/Mobile App high-value claims from new customers.",
        "Maintain a standing alert on gold_fraud_by_region for rate-without-volume anomalies.",
        "Feed dispositions back as labels to move from rules toward a supervised fraud model.",
        "Review the false-positive population to tune thresholds and reduce genuine-customer friction.",
    ]))

    e.append(Spacer(1, 16))
    e.append(hr())
    sig = Table([
        [Paragraph("<b>Prepared by</b><br/>J. Harrington<br/>Senior Investigator, SIU", S["Cell"]),
         Paragraph("<b>Reviewed by</b><br/>D. Osei<br/>Head of Claims Fraud", S["Cell"]),
         Paragraph("<b>Date</b><br/>17 January 2025", S["Cell"])],
    ], colWidths=[5.5 * cm, 5.5 * cm, 5.5 * cm])
    sig.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 4)]))
    e.append(sig)

    doc.build(e)
    print("wrote", path)


# ================================================================ COMPLIANCE
def build_compliance():
    path = os.path.join(OUT, "eu_compliance_policy.pdf")
    doc = make_doc(path, "Allianz Group",
                   "Data Protection & Compliance Office", "INTERNAL — COMPLIANCE POLICY")
    e = []
    e.append(Spacer(1, 6))
    e.append(Paragraph("INTERNAL — COMPLIANCE POLICY", S["Cls"]))
    e.append(Paragraph("EU Regulatory &amp; Compliance Framework for Insurance Fraud Analytics",
                       S["DocTitle"]))
    e.append(Paragraph("Obligations governing the use of personal data and automated analytics "
                       "to detect insurance fraud across the EU/EEA", S["SubTitle"]))
    e.append(control_table([
        ("Document reference", "COMP-EU-FRAUD-2024"),
        ("Version", "1.0 (approved)"),
        ("Owner", "Group Data Protection Office"),
        ("Approver", "Group Chief Compliance Officer"),
        ("Applies to", "Automated and analyst-assisted fraud detection, EU/EEA"),
        ("Effective date", "01 February 2025"),
        ("Next review", "01 February 2026"),
    ]))
    e.append(Spacer(1, 8))
    e.append(Paragraph("Synthetic teaching material for a workshop — not legal advice. "
                       "Deliberately generalised.", S["Small"]))

    e.append(Paragraph("1.  Purpose and Scope", S["H1"]))
    e.append(hr())
    e.append(Paragraph(
        "This framework sets out the regulatory obligations and internal controls that apply "
        "when the firm uses personal data and automated analytics to detect, investigate and "
        "act upon suspected insurance fraud within the European Union and wider EEA. It covers "
        "the full lifecycle implemented in this workshop: ingestion into <b>Bronze</b>; cleaning "
        "and enrichment into <b>Silver</b> and <b>Gold</b> (including <b>fraud_risk_score</b> and "
        "the four risk flags); natural-language querying via Genie; and analyst decisions to pay, "
        "deny or refer a claim.", S["Body2"]))
    e.append(Paragraph(
        "The dataset contains both non-personal claim attributes (policy_type, claim_amount, "
        "region) and <b>personal data</b> relating to identifiable individuals (policyholder_id, "
        "full_name, date_of_birth, email, phone, credit_score), bringing the activity within the "
        "GDPR.", S["Body2"]))

    e.append(Paragraph("2.  Governing Instruments", S["H1"]))
    e.append(hr())
    e.append(Paragraph("2.1  General Data Protection Regulation (Regulation (EU) 2016/679)", S["H2"]))
    e.append(Paragraph(
        "The GDPR is the central instrument. Fraud prevention is expressly recognised as a "
        "legitimate interest (Recital 47). The firm typically relies on <b>Article 6(1)(f) "
        "legitimate interests</b>, supported by a documented Legitimate Interests Assessment, "
        "and <b>Article 6(1)(c) legal obligation</b> where anti-fraud law compels processing. "
        "Where special-category data is involved (Article 9), e.g. health data in a Health "
        "claim, an additional condition is required.", S["Body2"]))
    e.append(Paragraph("2.2  Article 22 — Automated Decision-Making", S["H2"]))
    e.append(Paragraph(
        "Article 22 grants the right not to be subject to a decision based <i>solely</i> on "
        "automated processing producing legal or similarly significant effects — and denying a "
        "claim is such an effect. <b>Control implication:</b> the fraud_risk_score is used only "
        "to <b>triage and rank</b> claims for human review; it must never auto-deny a claim. A "
        "qualified analyst makes the final disposition, preserving the safeguards in Article "
        "22(2)–(3).", S["Body2"]))
    e.append(Paragraph("2.3  EU AI Act (Regulation (EU) 2024/1689)", S["H2"]))
    e.append(Paragraph(
        "AI systems used for risk assessment and pricing in life and health insurance are "
        "treated as high-risk, attracting obligations on data quality, logging, transparency, "
        "human oversight and accuracy. A fraud-triage model materially affecting access to an "
        "insurance benefit should be assessed against these obligations; the Act's principles "
        "are sound practice even where a system is not formally high-risk.", S["Body2"]))
    e.append(Paragraph("2.4  Sector and financial-crime rules", S["H2"]))
    e.append(bullets([
        "Insurance Distribution Directive (Directive (EU) 2016/97) — conduct and "
        "customer-treatment obligations constraining how fraud suspicion is acted upon.",
        "Anti-Money-Laundering framework (Directive (EU) 2015/849 as amended) — suspicious "
        "activity reporting where fraud intersects with money laundering.",
    ]))

    e.append(PageBreak())
    e.append(Paragraph("3.  Data-Protection Principles Applied", S["H1"]))
    e.append(hr())
    e.append(bullets([
        "<b>Lawfulness, fairness, transparency</b> — the privacy notice must disclose fraud "
        "processing; analytics must not discriminate against protected groups.",
        "<b>Purpose limitation</b> — data collected for underwriting/claims may be reused for "
        "fraud detection (compatible, disclosed) but not silently repurposed for marketing.",
        "<b>Data minimisation</b> — the Silver layer deliberately drops direct-contact PII "
        "(email, phone), retaining only region, tenure, credit score and age.",
        "<b>Accuracy</b> — data-quality expectations (unique non-null claim_id, claim_amount &gt; 0, "
        "report_date ≥ claim_date, is_fraud ∈ {0,1}) support this principle directly.",
        "<b>Storage limitation</b> — investigation records retained only as long as necessary, "
        "then deleted or anonymised.",
        "<b>Integrity &amp; confidentiality</b> — Unity Catalog enforces table/column access; the "
        "app runs under a least-privilege service principal; all access is auditable.",
        "<b>Accountability</b> — demonstrated via this framework, the LIA, a DPIA, Unity Catalog "
        "lineage and an audit trail of analyst decisions.",
    ]))

    e.append(Paragraph("4.  Mapping Data Fields to Compliance Categories", S["H1"]))
    e.append(hr())
    e.append(data_table(
        ["Data field", "Category", "Notes"],
        [["policyholder_id", "Pseudonymous ID", "Personal data when linkable to a person"],
         ["full_name, email, phone", "Direct PII", "Minimised — dropped from Silver policyholders"],
         ["date_of_birth / customer_age", "Personal data", "Age can proxy protected traits — monitor"],
         ["credit_score", "Personal (financial)", "Sensitive; drives low_credit_flag"],
         ["region", "Personal (location)", "Watch for geographic proxy discrimination"],
         ["gender", "Special-category-adjacent", "Must not drive fraud decisions"],
         ["claim_amount, policy_type", "Non-personal", "Low risk"],
         ["fraud_risk_score, is_fraud", "Derived / label", "Governed by Article 22 safeguards"]],
        [4.4 * cm, 3.6 * cm, 8.5 * cm]))

    e.append(Paragraph("5.  Required Controls", S["H1"]))
    e.append(hr())
    e.append(bullets([
        "Current DPIA (Article 35) and documented Legitimate Interests Assessment.",
        "Record of Processing Activities (Article 30).",
        "Human oversight: automated scores triage but never auto-deny.",
        "Fairness testing across protected characteristics to detect disparate impact.",
        "Least-privilege access via Unity Catalog; PII masking where feasible; full audit logging.",
        "Explainability: the additive, named-flag fraud_risk_score is preferred over an opaque model.",
    ]))

    e.append(Paragraph("6.  Data Subject Rights &amp; Breach Response", S["H1"]))
    e.append(hr())
    e.append(Paragraph(
        "Data subjects retain rights of information, access, rectification, objection and the "
        "Article 22 safeguards. Access may be lawfully restricted where disclosure would tip off "
        "a fraudster, provided the restriction is necessary, proportionate and recorded. Any "
        "personal-data breach must be assessed within the GDPR 72-hour notification window "
        "(Article 33) and, where high risk arises, affected individuals notified (Article 34).",
        S["Body2"]))

    e.append(Spacer(1, 14))
    e.append(hr())
    e.append(Paragraph("Approved by the Group Chief Compliance Officer. This document is "
                       "reviewed annually or upon material regulatory change.", S["Small"]))

    doc.build(e)
    print("wrote", path)


# ================================================================ GENIE DOCX
def build_genie_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    path = os.path.join(OUT, "genie_questions.docx")
    d = Document()

    styles_ = d.styles["Normal"]
    styles_.font.name = "Calibri"
    styles_.font.size = Pt(11)

    title = d.add_heading("Genie Questions — Insurance Fraud Analytics", level=0)
    sub = d.add_paragraph("A ready-to-use set of natural-language questions for the Insurance "
                          "Fraud Analytics Genie Space and the “Ask Genie” panel in the app. "
                          "Grouped from simple to advanced.")
    sub.runs[0].italic = True

    def section(title_, qs):
        d.add_heading(title_, level=1)
        for q in qs:
            p = d.add_paragraph(style="List Number")
            p.add_run(q)

    section("A. Warm-up (single metric)", [
        "What is the overall fraud rate?",
        "How many claims are in the dataset in total?",
        "What is the total payout across all claims, and how much of it is on fraudulent claims?",
    ])
    section("B. Breakdowns", [
        "Which regions have the highest fraud rate?",
        "Show the fraud rate by policy type, from highest to lowest.",
        "Which channel has the most fraudulent claims?",
        "What is the average claim amount for fraudulent versus legitimate claims?",
    ])
    section("C. Risk indicators", [
        "How many high-risk claims are there, where the fraud risk score is 3 or higher?",
        "What is the average fraud risk score for fraudulent claims compared with legitimate ones?",
        "List the 10 highest-value fraudulent claims with their region and policy type.",
        "How many claims were reported within 2 days of the incident and were also fraudulent?",
    ])
    section("D. Trends", [
        "Show the monthly fraud rate over time.",
        "Which month had the highest fraudulent payout?",
    ])
    section("E. Combined / analytical", [
        "For Auto claims in London submitted online, what is the fraud rate and how many claims are there?",
        "Which policy type and region combination has the worst fraud rate?",
        "Among new customers with low credit scores, how many claims were fraudulent?",
    ])

    d.add_heading("F. Questions to ask after uploading a document", level=1)
    d.add_paragraph("Upload a file in the “Ask Genie” panel (stored in raw/input/userdata), "
                    "then ask questions that combine the document with the data.")
    d.add_paragraph("After uploading the SIU fraud investigation report:", style="Intense Quote")
    for q in [
        "Based on the uploaded investigation report, which regions and channels should I focus on, "
        "and what does the data show the fraud rate is for those segments?",
        "The report describes a four-signal fingerprint. Show me claims that match all four signals "
        "(fraud risk score of 4) and how many were confirmed fraud.",
        "The report mentions an Auto theft reported immediately. How many Auto theft claims were "
        "reported within 1 day?",
    ]:
        p = d.add_paragraph(style="List Number")
        p.add_run(q)
    d.add_paragraph("After uploading the EU compliance framework:", style="Intense Quote")
    for q in [
        "According to the uploaded compliance framework, which fields in this dataset are personal "
        "data under GDPR?",
        "The compliance document says automated scores must not auto-deny claims. Explain how the "
        "fraud risk score is used for triage rather than automated denial.",
        "Which data fields does the compliance framework say should be dropped for data "
        "minimisation, and are they present in the gold table?",
    ]:
        p = d.add_paragraph(style="List Number")
        p.add_run(q)

    d.add_heading("Notes for facilitators", level=1)
    for n in [
        "If Genie struggles with a question, add it as an example query in the Genie Space "
        "instructions with the correct SQL.",
        "Questions in sections A–E make a good benchmark set for the Genie Benchmarks tab.",
        "Document-grounded questions (section F) rely on the app’s “Ask Genie” panel, which "
        "passes the uploaded document text as additional context.",
    ]:
        d.add_paragraph(n, style="List Bullet")

    d.save(path)
    print("wrote", path)


if __name__ == "__main__":
    build_fraud_report()
    build_compliance()
    build_genie_docx()
    print("DONE")
